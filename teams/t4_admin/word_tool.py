import json
import logging
from datetime import datetime
from pathlib import Path
from crewai.tools import BaseTool
from config import OUTPUT_DIR
from teams.t4_admin.styles import resolve_style, _hex_from_rgb, DEFAULT_STYLE_NAME
from teams.t4_admin.chart_utils import render_chart

logger = logging.getLogger('bima_core')


# ============================================================
# Word Generator — style + image embedding
# ============================================================
class WordGeneratorTool(BaseTool):
    name: str = "Word Generator Tool"
    description: str = """Buat file Word (.docx) dengan style fleksibel.
    Input format JSON string:
    {
        "filename": "nama_file",
        "style": "formal" | "semi_formal" | "informal" | "akademik",
        "title": "Judul Dokumen",
        "subtitle": "Subjudul opsional",
        "author": "Nama Author",
        "toc": true,
        "sections": [
            {
                "heading": "Judul Bagian",
                "content": "Isi paragraf...",
                "list": ["Item 1", "Item 2"],
                "key_values": {"Nama": "Bima", "Jabatan": "Admin"},
                "image_path": "/path/to/image.png",
                "charts": [
                    {"type": "bar", "title": "Penjualan 2026", "labels": ["Q1","Q2","Q3"],
                     "datasets": [{"label": "Unit", "data": [100, 150, 200]}]}
                ],
                "table": {
                    "headers": ["Kolom1", "Kolom2"],
                    "rows": [["data1", "data2"]]
                }
            }
        ],
        "references": [
            {"text": "Penulis (2026). Judul paper. Penerbit.",
             "url": "https://doi.org/..."}
        ]
    }
    Field 'charts' (opsional, list): tipe 'bar'|'line'|'pie' — auto-render via matplotlib pakai warna dari style.
    Field 'references' (opsional): daftar pustaka — auto render section terakhir dengan hyperlink clickable.
    URL referensi WAJIB valid (Wikipedia/situs resmi/jurnal open-access). JANGAN dikarang.
    Tipe dokumen support: laporan, proposal, surat, resume, blog, tutorial — gaya menyesuaikan field 'style'."""

    def _run(self, input_json: str) -> str:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml

            def set_cell_bg(cell, hex_color: str):
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
                cell._tc.get_or_add_tcPr().append(shd)

            def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
                """Atur padding cell (twips, 1 inch = 1440 twips, ~80 = 0.055 inch)."""
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_mar = OxmlElement('w:tcMar')
                for direction, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
                    node = OxmlElement(f'w:{direction}')
                    node.set(qn('w:w'), str(val))
                    node.set(qn('w:type'), 'dxa')
                    tc_mar.append(node)
                tc_pr.append(tc_mar)

            def add_hyperlink(paragraph, text: str, url: str, color_hex: str = "0563C1"):
                """Tambah hyperlink clickable ke paragraph (python-docx tidak punya API langsung)."""
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                part = paragraph.part
                r_id = part.relate_to(
                    url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                color_node = OxmlElement('w:color')
                color_node.set(qn('w:val'), color_hex)
                rPr.append(color_node)
                u_node = OxmlElement('w:u')
                u_node.set(qn('w:val'), 'single')
                rPr.append(u_node)
                new_run.append(rPr)
                t = OxmlElement('w:t')
                t.text = text
                t.set(qn('xml:space'), 'preserve')
                new_run.append(t)
                hyperlink.append(new_run)
                paragraph._p.append(hyperlink)
                return hyperlink

            try:
                data = json.loads(input_json)
            except json.JSONDecodeError as e:
                return f"FAILED|JSON tidak valid: {e}"

            style_name = data.get("style", DEFAULT_STYLE_NAME)
            style = resolve_style(data)
            title_color = RGBColor(*style["title_rgb"])
            accent_color = RGBColor(*style["accent_rgb"])
            header_hex = _hex_from_rgb(style["table_header_rgb"])

            # === Configurable typography (JSON overrides > style defaults) ===
            font_family = data.get("font_family", style.get("font_family", "Calibri"))
            margins_cm = data.get("margins", style.get("margins_cm", {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54}))
            do_justify = data.get("justify", style.get("justify", False))
            line_spacing_mult = data.get("line_spacing", style.get("line_spacing", 1.15))

            doc = Document()

            # === MARGIN — configurable via JSON atau style preset ===
            from docx.shared import Cm
            section_doc = doc.sections[0]
            # Min top_margin 2.0cm supaya title gak nabrak header area
            _top_cm = max(float(margins_cm.get("top", 2.54)), 2.0)
            section_doc.top_margin    = Cm(_top_cm)
            section_doc.bottom_margin = Cm(margins_cm.get("bottom", 2.54))
            section_doc.left_margin   = Cm(margins_cm.get("left", 2.54))
            section_doc.right_margin  = Cm(margins_cm.get("right", 2.54))
            # Explicit header/footer distance — default Word ~1.27cm, pendekkin ke 0.8cm
            # supaya body area gak ke-"makan" header
            section_doc.header_distance = Cm(0.8)
            section_doc.footer_distance = Cm(0.8)

            # === GAYA NORMAL: font, line spacing, space after ===
            normal_style = doc.styles['Normal']
            normal_style.font.name = font_family
            normal_style.font.size = Pt(style["body_size"])
            normal_style.paragraph_format.line_spacing = Pt(style["body_size"] * line_spacing_mult)
            normal_style.paragraph_format.space_after = Pt(6)

            # === HEADING STYLES (multi-level support) ===
            _heading_configs = [
                (1, style["heading_size"] + 2, True, False),   # Level 1: besar, bold
                (2, style["heading_size"], True, False),       # Level 2: sedang, bold
                (3, style["heading_size"] - 1, True, True),    # Level 3: kecil, bold italic
            ]
            for _lvl, _sz, _bold, _italic in _heading_configs:
                _hstyle_name = f'Heading {_lvl}'
                if _hstyle_name in doc.styles:
                    _hs = doc.styles[_hstyle_name]
                    _hs.font.name = font_family
                    _hs.font.size = Pt(_sz)
                    _hs.font.bold = _bold
                    _hs.font.italic = _italic
                    _hs.font.color.rgb = title_color
                    _hs.paragraph_format.space_before = Pt(12)
                    _hs.paragraph_format.space_after = Pt(6)

            # === HELPER: Add section break + set page number format ===
            def _set_page_num_format(section_obj, fmt='decimal', start=None):
                """Set page number format. fmt: 'decimal', 'lowerRoman', 'upperRoman'."""
                from docx.oxml import OxmlElement as _OE
                from docx.oxml.ns import qn as _qn
                sectPr = section_obj._sectPr
                pgNumType = _OE('w:pgNumType')
                pgNumType.set(_qn('w:fmt'), fmt)
                if start is not None:
                    pgNumType.set(_qn('w:start'), str(start))
                sectPr.append(pgNumType)

            # === PAGE NUMBERS DI FOOTER ===
            footer = section_doc.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run()
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

            # Tambah field PAGE dan NUMPAGES via XML
            def _add_fld(para, instr):
                from docx.oxml import OxmlElement as _OE
                from docx.oxml.ns import qn as _qn
                fldChar_b = _OE('w:fldChar'); fldChar_b.set(_qn('w:fldCharType'), 'begin')
                instrText = _OE('w:instrText'); instrText.text = instr
                instrText.set(_qn('xml:space'), 'preserve')
                fldChar_e = _OE('w:fldChar'); fldChar_e.set(_qn('w:fldCharType'), 'end')
                r = _OE('w:r')
                r.append(fldChar_b)
                p = _OE('w:r')
                p.append(instrText)
                e = _OE('w:r')
                e.append(fldChar_e)
                para._p.extend([r, p, e])
            footer_para.add_run(f"{data.get('author', 'B.I.M.A Core')}  |  halaman ")
            _add_fld(footer_para, ' PAGE ')
            footer_para.add_run(" dari ")
            _add_fld(footer_para, ' NUMPAGES ')

            # === Set front matter to Roman numerals ===
            is_akademik = (style_name == 'akademik')
            if is_akademik:
                _set_page_num_format(section_doc, fmt='lowerRoman', start=1)

            # Title
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(18)
            title_para.paragraph_format.space_after = Pt(6)
            title_para.paragraph_format.keep_with_next = True
            title_run = title_para.add_run(data.get("title", "Dokumen"))
            title_run.font.bold = True
            title_run.font.size = Pt(style["title_size"] + 4)
            title_run.font.color.rgb = title_color

            # Subtitle (opsional)
            if data.get("subtitle"):
                sub_para = doc.add_paragraph()
                sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_run = sub_para.add_run(data["subtitle"])
                sub_run.font.italic = True
                sub_run.font.size = Pt(style["body_size"] + 2)
                sub_run.font.color.rgb = accent_color

            # Author + tanggal
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.add_run(f"{data.get('author', 'B.I.M.A Core')} · {datetime.now().strftime('%d %B %Y')} · [{style['label']}]")
            info_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            info_run.font.size = Pt(9)
            doc.add_paragraph()

            # === ABSTRAK (opsional — khusus dokumen akademik) ===
            if data.get("abstract"):
                doc.add_page_break()
                abs_title = doc.add_paragraph()
                abs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                abs_title_run = abs_title.add_run("ABSTRAK")
                abs_title_run.font.bold = True
                abs_title_run.font.size = Pt(style["heading_size"] + 2)
                abs_title_run.font.color.rgb = title_color
                abs_title_run.font.name = font_family

                abs_para = doc.add_paragraph(data["abstract"])
                abs_para.paragraph_format.space_after = Pt(8)
                # Abstrak selalu single-spaced meskipun body 1.5
                abs_para.paragraph_format.line_spacing = Pt(style["body_size"] * 1.15)
                abs_para.paragraph_format.left_indent = Inches(0.5)
                abs_para.paragraph_format.right_indent = Inches(0.5)
                if do_justify:
                    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Keywords
                if data.get("keywords"):
                    kw_para = doc.add_paragraph()
                    kw_para.paragraph_format.left_indent = Inches(0.5)
                    kw_label = kw_para.add_run("Kata Kunci: ")
                    kw_label.font.bold = True
                    kw_label.font.size = Pt(style["body_size"])
                    kw_label.font.name = font_family
                    kw_value = kw_para.add_run(", ".join(data["keywords"]))
                    kw_value.font.italic = True
                    kw_value.font.size = Pt(style["body_size"])
                    kw_value.font.name = font_family

            # === Table of Contents (opsional, multi-level) ===
            sections = data.get("sections", [])
            if data.get("toc") and sections:
                doc.add_page_break()
                toc_h = doc.add_paragraph()
                toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                toc_run = toc_h.add_run("DAFTAR ISI")
                toc_run.font.bold = True
                toc_run.font.size = Pt(style["heading_size"] + 2)
                toc_run.font.color.rgb = title_color
                toc_run.font.name = font_family
                doc.add_paragraph()

                for sec in sections:
                    if sec.get("heading"):
                        lvl = sec.get("level", 1)
                        toc_p = doc.add_paragraph()
                        toc_p.paragraph_format.space_after = Pt(2)
                        # Indentasi berdasarkan level
                        toc_p.paragraph_format.left_indent = Inches(0.3 * (lvl - 1))
                        toc_run = toc_p.add_run(sec["heading"])
                        toc_run.font.size = Pt(style["body_size"])
                        toc_run.font.name = font_family
                        if lvl == 1:
                            toc_run.font.bold = True

            # === SECTION BREAK: Roman → Arabic page numbering ===
            if is_akademik:
                from docx.oxml import OxmlElement as _OE_sb
                from docx.oxml.ns import qn as _qn_sb
                # Add section break (new page) before body content
                body_break_para = doc.add_paragraph()
                pPr = body_break_para._p.get_or_add_pPr()
                sectPr = _OE_sb('w:sectPr')
                sectType = _OE_sb('w:type')
                sectType.set(_qn_sb('w:val'), 'nextPage')
                sectPr.append(sectType)
                # Copy margins to new section
                pgMar = _OE_sb('w:pgMar')
                pgMar.set(_qn_sb('w:top'), str(int(margins_cm.get('top', 2.54) * 567)))
                pgMar.set(_qn_sb('w:bottom'), str(int(margins_cm.get('bottom', 2.54) * 567)))
                pgMar.set(_qn_sb('w:left'), str(int(margins_cm.get('left', 2.54) * 567)))
                pgMar.set(_qn_sb('w:right'), str(int(margins_cm.get('right', 2.54) * 567)))
                sectPr.append(pgMar)
                # Set Arabic numbering starting from 1
                pgNumType = _OE_sb('w:pgNumType')
                pgNumType.set(_qn_sb('w:fmt'), 'decimal')
                pgNumType.set(_qn_sb('w:start'), '1')
                sectPr.append(pgNumType)
                pPr.append(sectPr)
            else:
                if data.get("toc") and sections:
                    doc.add_page_break()

            for section in sections:
                if section.get("heading"):
                    lvl = section.get("level", 1)
                    lvl = max(1, min(lvl, 3))  # clamp 1-3

                    # Gunakan built-in Heading style (sudah dikustomisasi di atas)
                    h = doc.add_heading(section["heading"], level=lvl)
                    # Override font untuk memastikan konsistensi
                    for run in h.runs:
                        run.font.name = font_family
                        run.font.color.rgb = title_color

                    # Garis aksen hanya untuk heading level 1
                    if lvl == 1:
                        h.paragraph_format.space_after = Pt(4)
                        from docx.oxml import OxmlElement as _OE2
                        from docx.oxml.ns import qn as _qn2
                        pPr = h._p.get_or_add_pPr()
                        pBdr = _OE2('w:pBdr')
                        bottom = _OE2('w:bottom')
                        bottom.set(_qn2('w:val'), 'single')
                        bottom.set(_qn2('w:sz'), '4')
                        bottom.set(_qn2('w:space'), '4')
                        bottom.set(_qn2('w:color'), _hex_from_rgb(style["accent_rgb"]))
                        pBdr.append(bottom)
                        pPr.append(pBdr)
                    elif lvl == 2:
                        h.paragraph_format.space_after = Pt(4)
                        h.paragraph_format.left_indent = Inches(0.15)
                    elif lvl == 3:
                        h.paragraph_format.space_after = Pt(3)
                        h.paragraph_format.left_indent = Inches(0.3)

                if section.get("content"):
                    p = doc.add_paragraph(section["content"])
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.line_spacing = Pt(style["body_size"] * line_spacing_mult)
                    if do_justify:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Image embedding
                if section.get("image_path"):
                    img_path = Path(section["image_path"])
                    if img_path.exists() and img_path.is_file():
                        try:
                            doc.add_picture(str(img_path), width=Inches(5.5))
                        except Exception as img_err:
                            logger.warning(f"[ADMIN] Gagal embed image {img_path}: {img_err}")

                # Chart rendering (matplotlib → embed PNG)
                for chart in section.get("charts", []) or []:
                    try:
                        chart_path = render_chart(chart, style)
                        doc.add_picture(chart_path, width=Inches(6))
                    except Exception as chart_err:
                        logger.warning(f"[ADMIN] Gagal render chart Word: {chart_err}")

                # Bullet list
                if section.get("list"):
                    for item in section["list"]:
                        doc.add_paragraph(str(item), style='List Bullet')

                # Key Values (untuk surat izin, detail rapi dengan titik dua sejajar)
                if section.get("key_values") and isinstance(section["key_values"], dict):
                    kv_table = doc.add_table(rows=len(section["key_values"]), cols=3)
                    for idx, (k, v) in enumerate(section["key_values"].items()):
                        row = kv_table.rows[idx].cells
                        row[0].text = str(k)
                        row[0].width = Inches(1.5)
                        row[1].text = ":"
                        row[1].width = Inches(0.2)
                        row[2].text = str(v)
                        for cell in row:
                            set_cell_margins(cell, top=20, bottom=20, left=40, right=40)
                    doc.add_paragraph()

                # Table
                if section.get("table"):
                    tbl_data = section["table"]
                    headers = tbl_data.get("headers", [])
                    rows = tbl_data.get("rows", [])
                    if headers and rows:
                        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                        table.style = 'Table Grid'
                        table.autofit = True

                        hdr_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            hdr_cells[i].text = ""
                            set_cell_bg(hdr_cells[i], header_hex)
                            set_cell_margins(hdr_cells[i], top=100, bottom=100, left=140, right=140)
                            para = hdr_cells[i].paragraphs[0]
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run(str(header))
                            run.font.bold = True
                            run.font.size = Pt(style["body_size"])
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                        for row_idx, row_data in enumerate(rows, 1):
                            row_cells = table.rows[row_idx].cells
                            for col_idx, value in enumerate(row_data):
                                if col_idx < len(row_cells):
                                    row_cells[col_idx].text = str(value)
                                    set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=140, right=140)
                                    if row_idx % 2 == 0:
                                        set_cell_bg(row_cells[col_idx], _hex_from_rgb(style["table_alt_rgb"]))

                doc.add_paragraph()

            # === DAFTAR PUSTAKA ===
            references = data.get("references", [])
            if references:
                doc.add_paragraph()
                ref_h = doc.add_paragraph()
                ref_run = ref_h.add_run("Daftar Pustaka")
                ref_run.font.bold = True
                ref_run.font.size = Pt(style["heading_size"])
                ref_run.font.color.rgb = title_color

                accent_hex = _hex_from_rgb(style["accent_rgb"])
                for idx, ref in enumerate(references, 1):
                    if isinstance(ref, dict):
                        ref_text = ref.get("text", "")
                        ref_url = ref.get("url", "")
                    else:
                        ref_text = str(ref)
                        ref_url = ""

                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.left_indent = Inches(0.25)
                    main_run = p.add_run(f"{idx}. {ref_text}")
                    main_run.font.size = Pt(style["body_size"])

                    if ref_url:
                        p.add_run("  ").font.size = Pt(style["body_size"])
                        add_hyperlink(p, ref_url, ref_url, color_hex=accent_hex)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{data.get('filename', 'dokumen')}_{timestamp}.docx"
            filepath = OUTPUT_DIR / filename
            doc.save(filepath)

            return f"SUCCESS|{filepath}|Word ({style['label']}) berhasil dibuat: {filename}"
        except Exception as e:
            logger.error(f"[ADMIN] WordGenerator error: {e}", exc_info=True)
            return f"FAILED|{e}"
